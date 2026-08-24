#!/usr/bin/env python3
"""将 markdown 需求文档转换为带样式的 docx。

用法:
    python build_requirement_docx.py                              # 默认生成 V1.8 docx
    python build_requirement_docx.py --source X.md --output Y.docx # 自定义输入输出

特性:
    - markdown 标题/列表/表格/内联代码渲染
    - ASCII 线框代码块自动渲染为窗口风格 PNG 图片（保证对齐）
    - 封面 + 自动目录 + 页眉页脚
"""
from __future__ import annotations

import io
import random
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont, ImageFilter

    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "requirement.md"
OUTPUT = ROOT / "仿真数模资产管理系统_产品需求文档_V1.8.docx"

BODY_FONT = "PingFang SC"
HEADING_FONT = "PingFang SC"
MONO_FONT = "Consolas"
ACCENT = "1F4E78"
LIGHT_ACCENT = "D9EAF7"
LIGHT_GRAY = "F2F2F2"

# ── ASCII 线框渲染配置（Balsamiq 手绘风） ──────────────────────────
HAND_EN_FONT = "/System/Library/Fonts/Supplemental/MarkerFelt.ttc"
HAND_CN_FONT = "/System/Library/Fonts/Supplemental/Comic Sans MS.ttf"
# 转角字符 = 真线框特征；竖线/箭头 = 状态机或流程图（保持文本）
WIREFRAME_MARKERS = ("┌", "└", "├", "┐", "┘", "┤")
# 线框结构字符：横线/竖线延续（用于解析网格画手绘线）
H_CHARS = set("─├┤┌┐")
V_CHARS = set("│├┤┌└")
TEXT_SKIP = set("│─┌┐└┘├┤")


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_inline(paragraph, text: str):
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, BODY_FONT)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, BODY_FONT, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, MONO_FONT)
            run.font.color.rgb = RGBColor(80, 80, 80)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, BODY_FONT)


# ── ASCII 线框 → Balsamiq 手绘风 PNG ────────────────────────────────
def _is_wide(ch: str) -> bool:
    if ch in TEXT_SKIP:
        return False
    return ord(ch) >= 0x2E80


def _draw_sketchy_line(draw, x1, y1, x2, y2, color, width=3, seed=7, orientation="h"):
    """手绘抖动线：横线轻微纵向抖、竖线轻微横向抖，保持方向感。"""
    rng = random.Random(seed)
    length = abs((x2 - x1) if orientation == "h" else (y2 - y1))
    n = max(5, int(length / 4))
    prev = None
    for i in range(n + 1):
        t = i / n
        if orientation == "h":
            x = x1 + (x2 - x1) * t
            y = y1 + rng.uniform(-0.45, 0.45)
        else:
            y = y1 + (y2 - y1) * t
            x = x1 + rng.uniform(-0.4, 0.4)
        pt = (x, y)
        if prev is not None:
            draw.line([prev, pt], fill=color, width=width)
        prev = pt


def _render_wireframe_png(code_lines: list[str]) -> bytes | None:
    """把 ASCII 线框渲染为 Balsamiq 手绘风 PNG（手写字体 + 抖动线条 + 米黄画布）。

    解析线框网格结构：横线/竖线分别用方向感知抖动手绘，文字用 Marker Felt /
    Comic Sans（中文）渲染。返回 PNG bytes；缺 Pillow 或字体时返回 None。
    """
    if not _HAS_PIL:
        return None
    try:
        hand_en = ImageFont.truetype(HAND_EN_FONT, 15)
        hand_cn = ImageFont.truetype(HAND_CN_FONT, 15)
    except Exception:
        return None

    cell_w = hand_en.getbbox("M")[2]
    asc, desc = hand_en.getmetrics()
    line_h = asc + desc + 6
    pad_x, pad_y = 16, 14
    fg = "#2D2D2D"
    canvas = "#FFFDF4"
    border = "#4A4A4A"

    # 画布宽度按字符累加像素宽计算（全角=2 单元）
    content_w = max(sum(cell_w * (2 if _is_wide(ch) else 1) for ch in line) for line in code_lines) + pad_x * 2
    content_h = len(code_lines) * line_h + pad_y * 2
    img = Image.new("RGB", (content_w, content_h), canvas)
    draw = ImageDraw.Draw(img)

    def row_y(r):
        return pad_y + r * line_h + line_h / 2

    # 预计算每行"列 → 像素 x"映射（全角字符占 2 单元）
    col_px = {}
    for r, line in enumerate(code_lines):
        x = pad_x
        for c, ch in enumerate(line):
            col_px[(r, c)] = x + cell_w / 2
            x += cell_w * (2 if _is_wide(ch) else 1)

    # 横线（─ 连续段，含 ├┤┌┐ 端点）
    for r, line in enumerate(code_lines):
        seg_start = None
        for c, ch in enumerate(line):
            if ch in H_CHARS:
                if seg_start is None:
                    seg_start = c
            else:
                if seg_start is not None:
                    _draw_sketchy_line(
                        draw, col_px[(r, seg_start)], row_y(r), col_px[(r, c - 1)], row_y(r),
                        border, seed=r * 10 + c, orientation="h")
                    seg_start = None
        if seg_start is not None:
            _draw_sketchy_line(
                draw, col_px[(r, seg_start)], row_y(r), col_px[(r, len(line) - 1)], row_y(r),
                border, seed=r * 10 + len(line), orientation="h")

    # 竖线（│ 连续段，含 ├┤┌└ 端点）
    max_cols = max(len(l) for l in code_lines)
    for c in range(max_cols):
        seg_start = None
        for r in range(len(code_lines)):
            ch = code_lines[r][c] if c < len(code_lines[r]) else " "
            if ch in V_CHARS:
                if seg_start is None:
                    seg_start = r
            else:
                if seg_start is not None and (seg_start, c) in col_px and (r - 1, c) in col_px:
                    _draw_sketchy_line(
                        draw, col_px[(seg_start, c)], row_y(seg_start), col_px[(r - 1, c)], row_y(r - 1),
                        border, seed=c * 10 + seg_start, orientation="v")
                    seg_start = None
        if seg_start is not None and (seg_start, c) in col_px and (len(code_lines) - 1, c) in col_px:
            _draw_sketchy_line(
                draw, col_px[(seg_start, c)], row_y(seg_start), col_px[(len(code_lines) - 1, c)], row_y(len(code_lines) - 1),
                border, seed=c * 10 + seg_start, orientation="v")

    # 文字（跳过结构字符）
    y = pad_y - 2
    for line in code_lines:
        x = pad_x
        for ch in line:
            if ch in TEXT_SKIP:
                x += cell_w
                continue
            font = hand_cn if _is_wide(ch) else hand_en
            draw.text((x, y), ch, font=font, fill=fg)
            x += cell_w * (2 if _is_wide(ch) else 1)
        y += line_h

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Title": (HEADING_FONT, 24, ACCENT, 0, 18),
        "Heading 1": (HEADING_FONT, 16, ACCENT, 12, 8),
        "Heading 2": (HEADING_FONT, 14, ACCENT, 10, 6),
        "Heading 3": (HEADING_FONT, 12, "000000", 8, 4),
        "Heading 4": (HEADING_FONT, 11, "000000", 6, 3),
    }
    for style_name, (font, size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Requirement ID" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Requirement ID", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 4"]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("仿真数模资产管理系统产品需求文档")
    set_run_font(run, BODY_FONT, 9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("内部资料  |  第 ")
    set_run_font(run, BODY_FONT, 9)
    add_field(p, "PAGE")
    run = p.add_run(" 页")
    set_run_font(run, BODY_FONT, 9)


def add_cover(doc: Document, version: str = "V1.8", date: str = "2026-08-10"):
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("仿真数模资产管理系统").font.size = Pt(26)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("产品需求文档")
    set_run_font(run, HEADING_FONT, 20, True)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    for _ in range(5):
        doc.add_paragraph()
    details = [
        ("文档版本", version),
        ("文档状态", "产品需求基线"),
        ("适用终端", "PC Web"),
        ("编制日期", date),
    ]
    table = doc.add_table(rows=len(details) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.cell(0, 0).text = "文档属性"
    table.cell(0, 1).text = "内容"
    mark_header_row(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell, 120, 160, 120, 160)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, BODY_FONT, 11, True)
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row, (label, value) in zip(table.rows[1:], details):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 120, 160, 120, 160)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, BODY_FONT, 11, cell == row.cells[0])
        set_cell_shading(row.cells[0], LIGHT_ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("部门内部知识资产建设项目")
    set_run_font(run, BODY_FONT, 11)
    run.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_page_break()

    toc_heading = doc.add_paragraph("目录", style="Heading 1")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("在 Microsoft Word 中打开文档后更新目录域即可生成页码。")
    set_run_font(run, BODY_FONT, 9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_page_break()


def parse_table(lines: list[str], start: int):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    mark_header_row(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            add_inline(p, text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, ACCENT)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
            elif r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code_lines: list[str]):
    """渲染 markdown 代码块：ASCII 线框 → PNG 图片，其余 → 等宽文本。"""
    is_wireframe = False
    if code_lines:
        corner_lines = [l for l in code_lines if any(m in l for m in WIREFRAME_MARKERS)]
        first = code_lines[0].lstrip()
        last = code_lines[-1].lstrip()
        # 真线框：首行 ┌/└ 开头 + 末行 └ 结尾 + ≥2 行转角字符，排除状态机/流程图
        if (
            len(corner_lines) >= 2
            and (first.startswith("┌") or first.startswith("└"))
            and last.startswith("└")
        ):
            is_wireframe = True
    if is_wireframe:
        png = _render_wireframe_png(code_lines)
        if png is not None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(io.BytesIO(png))
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            return
    for code_line in code_lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(code_line if code_line else " ")
        set_run_font(run, MONO_FONT, 8.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
        if any(m in code_line for m in WIREFRAME_MARKERS):
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), LIGHT_GRAY)
            pPr.append(shd)


def add_body(doc: Document, markdown: str):
    lines = markdown.splitlines()
    index = 0
    skipped_source_title = False
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not skipped_source_title:
                skipped_source_title = True
                index += 1
                continue
            style = f"Heading {min(level, 4)}"
            p = doc.add_paragraph(style=style)
            add_inline(p, text)
            index += 1
            continue

        req = re.match(r"^\*\*([A-Z]+(?:-[A-Z]+)*-\d+\s+.+?)\*\*$", stripped)
        if req:
            p = doc.add_paragraph(style="Requirement ID")
            add_inline(p, req.group(1))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(1))
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines):
                code_line = lines[index].rstrip()
                if code_line.strip().startswith("```"):
                    index += 1
                    break
                code_lines.append(code_line)
                index += 1
            add_code_block(doc, code_lines)
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        index += 1


def set_update_fields(doc: Document):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main():
    import argparse

    parser = argparse.ArgumentParser(description="将 markdown 需求文档转换为 docx")
    parser.add_argument("--source", default=str(SOURCE), help="源 markdown 文件路径")
    parser.add_argument("--output", default=str(OUTPUT), help="输出 docx 文件路径")
    args = parser.parse_args()

    source_path = Path(args.source)
    output_path = Path(args.output)
    stem = source_path.stem

    markdown = source_path.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc, version="V2.0" if "v2" in stem.lower() else "V1.8")
    add_body(doc, markdown)
    set_update_fields(doc)

    doc.core_properties.title = f"仿真数模资产管理系统产品需求文档（{stem}）"
    doc.core_properties.subject = "仿真数模资产标准化治理、关联展示、产品功能与用户操作需求"
    doc.core_properties.author = "项目需求组"
    doc.core_properties.keywords = "图纸管理, 生产资料, 软件需求规格说明书, SRS"
    doc.core_properties.comments = f"面向产品、设计、开发、测试和外部评审的需求基线（{stem}）"

    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
